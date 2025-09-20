from io import BytesIO
try:
    from docx import Document
except Exception:
    Document = None

class DocxTextParser:
    def to_text(self, file_bytes: bytes) -> str:
        if Document is None:
            return file_bytes.decode("utf-8", "ignore")
        doc = Document(BytesIO(file_bytes))
        parts = []

        for p in doc.paragraphs:
            if p.text and p.text.strip():
                parts.append(p.text.strip())

        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        return "\n".join(parts)
